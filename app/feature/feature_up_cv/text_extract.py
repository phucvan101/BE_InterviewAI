from __future__ import annotations

import io
from pathlib import Path

import pypdf
from docx import Document as DocxDocument


class UnsupportedFileTypeError(Exception):
    pass


def extract_text_from_pdf(file_path: str) -> str:
    # Some PDFs may have leading whitespace/newlines before %PDF header.
    # pypdf can reject these ("invalid pdf header: b'\\n%PDF'"), so we normalize bytes.
    with open(file_path, "rb") as f:
        data = f.read()
    data = data.lstrip()  # remove leading whitespace/newlines
    with io.BytesIO(data) as bio:
        reader = pypdf.PdfReader(bio)
        chunks: list[str] = []
        for page in reader.pages:
            # Some PDFs return None for extract_text()
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs)
    # Include table content when available
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text_auto(file_path: str) -> str:
    """
    Extract text from .pdf, .docx, or .txt based on file suffix.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix == ".docx":
        return extract_text_from_docx(file_path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8").strip()
    raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")

