from __future__ import annotations

import io
import re
from pathlib import Path

import pypdf
from docx import Document as DocxDocument


class UnsupportedFileTypeError(Exception):
    pass


def extract_text_from_pdf(file_path: str) -> str:
    # Some PDFs may have leading whitespace/newlines before %PDF header.
    # pypdf can reject these ("invalid pdf header: b'\\n%PDF'"), so we normalize bytes.
    with open(file_path, "rb") as f:
        data = f.read()
    data = data.lstrip()  # remove leading whitespace/newlines
    with io.BytesIO(data) as bio:
        reader = pypdf.PdfReader(bio)
        chunks: list[str] = []
        for page in reader.pages:
            # Some PDFs return None for extract_text()
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs)
    # Include table content when available
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def clean_text_for_llm(text: str) -> str:
    """
    Remove emojis, icons, and non-standard characters from text.
    Keeps ASCII, Vietnamese characters, and common punctuation.
    """
    if not text:
        return text
    
    # Keep characters in range:
    # \u0000-\u1EFF : Basic Latin, Latin-1, Latin Extended (including Vietnamese)
    # \u2000-\u206F : General Punctuation (bullets, en-dash, em-dash, quotes)
    # \u20A0-\u20CF : Currency Symbols (₫, €, etc.)
    # \u2100-\u214F : Letterlike Symbols
    # \u25A0-\u25FF : Geometric Shapes (■, □, ▪, ▫)
    # \u2713\u2714 : Check marks (✓, ✔)
    cleaned = re.sub(r'[^\u0000-\u1EFF\u2000-\u206F\u20A0-\u20CF\u2100-\u214F\u25A0-\u25FF\u2713\u2714\n\r\t]', '', text)
    
    # Remove multiple spaces
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    # Remove multiple blank lines
    cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
    
    return cleaned.strip()


def extract_text_auto(file_path: str) -> str:
    """
    Extract text from .pdf, .docx, or .txt based on file suffix.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    text = ""
    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
    elif suffix == ".txt":
        text = path.read_text(encoding="utf-8").strip()
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")
        
    return clean_text_for_llm(text)

