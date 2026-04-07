# -*- coding: utf-8 -*-
"""
Extract company research information from documents
- Industry/Field
- Company Mission
- Values
- Culture
- Key achievements
- Skills/Technologies used
"""

import sys
import json
import re
from pathlib import Path
from docx import Document as DocxDocument
import pypdf
import google.generativeai as genai
import os
from dotenv import load_dotenv

load_dotenv()

GOOGLE_API_KEY = os.getenv('GEMINI_API_KEY')
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text()
        return text.strip()
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
        return text.strip()
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return ""


def extract_company_info(file_path: str) -> dict:
    """
    Extract company information using Gemini LLM
    
    Returns:
    {
        "company_name": "...",
        "industry": "...",
        "description": "...",
        "mission": "...",
        "values": [...],
        "key_skills": [...],
        "technologies": [...],
        "company_culture": "...",
        "key_achievements": [...]
    }
    """
    
    # ── Extract text ─────────────────────────────────
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        return {
            "success": False,
            "error": f"File not found: {file_path}"
        }
    
    filename_lower = str(file_path).lower()
    
    if filename_lower.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename_lower.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        return {
            "success": False,
            "error": "Unsupported file type"
        }
    
    if not text:
        return {
            "success": False,
            "error": "Could not extract text from file"
        }
    
    # ── LLM Extraction ───────────────────────────────
    try:
        prompt = f"""
Analyze this company document and extract key information in JSON format:

Document content:
{text[:3000]}  # Limit to first 3000 chars for API

Return ONLY valid JSON with these fields:
{{
    "company_name": "Company name (if found)",
    "industry": "Industry/Field (e.g., Technology, Finance, Healthcare)",
    "description": "Brief company description (2-3 sentences)",
    "mission": "Company mission statement",
    "values": ["value1", "value2", ...],
    "key_skills": ["skill1", "skill2", "skill3", ...],
    "technologies": ["Tech1", "Tech2", ...],
    "company_culture": "Brief description of company culture",
    "key_achievements": ["achievement1", "achievement2", ...]
}}

If information is not found for a field, use reasonable defaults or empty arrays.
Make sure response is VALID JSON only, no markdown or extra text.
"""
        
        print("\n🔄 Extracting company info with Gemini...")
        response = model.generate_content(prompt)
        
        if not response.text:
            return {
                "success": False,
                "error": "Gemini returned empty response"
            }
        
        # Parse JSON response
        response_text = response.text.strip()
        
        # Try to find JSON in response
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
            company_info = json.loads(json_str)
        else:
            company_info = json.loads(response_text)
        
        company_info["success"] = True
        return company_info
        
    except json.JSONDecodeError as e:
        print(f"❌ JSON parse error: {e}")
        print(f"   Response: {response.text[:200]}")
        return {
            "success": False,
            "error": f"Failed to parse Gemini response: {str(e)}"
        }
    except Exception as e:
        print(f"❌ Gemini extraction error: {e}")
        return {
            "success": False,
            "error": f"Gemini API error: {str(e)}"
        }


if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        result = extract_company_info(file_path)
        print(json.dumps(result, indent=2, ensure_ascii=False))
